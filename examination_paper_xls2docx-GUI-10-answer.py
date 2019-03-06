# coding=utf-8
from Tkinter import *
from tkFileDialog import *
from openpyxl import Workbook
from openpyxl import load_workbook
from openpyxl.styles import Border, Side, Font #设置字体和边框需要的模块
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH #段落居中
import random
import os
import sys
#reload(sys)
#sys.setdefaultencoding( "utf-8" )


#创建容器
tk=Tk()
tk.title("题库生成试卷的GUI")
mainfarm=Frame(tk,width=400, height=100,bg="green")
mainfarm.grid_propagate(0)
mainfarm.grid()
fram=Frame(mainfarm,width=400, height=100,bg="green")
fram.grid_propagate(0)
fram.grid()

e = Entry(fram)
e.grid(row=0,column=1)

e.delete(0, END)  # 将输入框里面的内容清空
e.insert(0, '显示文件路径')
filepath = StringVar()


def filefound():
    filepath = askopenfilename()
    print filepath
    e.delete(0, END)  # 将输入框里面的内容清空
    e.insert(0, filepath)
    #有时候我们希望读取到公式计算出来的结果，可以使用load_workbook()中的data_only属性
    #wb = load_workbook(u'F:\\PY\\Scripts\\Scripts\\题库（单选+多选+判断）.xlsx',data_only=True)
    wb = load_workbook(filepath, data_only=True)
    
    sheets = wb.get_sheet_names()
    print sheets
    dan_sheet = wb.get_sheet_by_name(sheets[0])
    duo_sheet = wb.get_sheet_by_name(sheets[1])
    panduan_sheet = wb.get_sheet_by_name(sheets[2])

    ws = wb.active
    rows = []
    for row in ws.iter_rows():
        rows.append(row)

    print u"行高：",ws.max_row
    print u"列宽：",ws.max_column

    duanhao = '、'
    huiche = '\n'
    
    for count in range(10):
        #打开文档
        document = Document()
        answerdocument = Document()
        
        #修改正文的中文字体类型，示例代码：（全局设置）
        document.styles['Normal'].font.name=u'仿宋'
        document.styles['Normal'].font.size =Pt(16) #16对应仿宋三号
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')

        head_str = u'题库随机抽题考试'+str(count)+u'\n' #文章标题
        # document.add_heading(head_str,0)
        head_paragraph = document.add_paragraph('')#添加一个段落
        head_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER #段落居中
        run = head_paragraph.add_run(head_str)
        run.bold = True   # 粗体是
        run.italic = False   # 斜体否


        document.styles['Normal'].font.size =Pt(12) #12对应仿宋小四号

        danxuanti_para = u'第一部分  单选题(20题)'
        dan_paragraph = document.add_paragraph(danxuanti_para)#考题
        answerdocument.add_paragraph(danxuanti_para) #答案
        
        dan_num=1
        dan_set = set()
        answer_dan_para = u''
        while(len(dan_set) < 20):
            i = random.randint(1, dan_sheet.max_row-2)
            if i not in dan_set:
               dan_set.add(i)
               dan_para = str(dan_num) + u'、' + dan_sheet.cell(row=i + 2, column=4).value + u'\n'  # 问题
               print dan_sheet.cell(row=i + 2, column=1).value,dan_sheet.cell(row=i + 2, column=4).value
               
               dan_temp = u'   A、%s\n   B、%s\n   C、%s\n   D、%s\n' % (
               dan_sheet.cell(row=i + 2, column=5).value, dan_sheet.cell(row=i + 2, column=6).value,
               dan_sheet.cell(row=i + 2, column=7).value, dan_sheet.cell(row=i + 2, column=8).value)
               dan_para += dan_temp
               print dan_para
               dan_paragraph = document.add_paragraph(dan_para)

               answer_dan_para += str(dan_num) + u'、' + dan_sheet.cell(row=i + 2, column=10).value + u'   '
               
               dan_num+=1               
            else:
                pass

        answerdocument.add_paragraph(answer_dan_para) #答案
        
        duoxuanti_para = u'第二部分  多选题(10题)'
        duo_paragraph = document.add_paragraph(duoxuanti_para)
        answerdocument.add_paragraph(duoxuanti_para) #答案
        
        duo_num=1
        duo_set = set()
        answer_duo_para = u''
        while(len(duo_set) < 10):
            i = random.randint(1, duo_sheet.max_row-2)
            if i not in duo_set:
               print i, duo_sheet.cell(row=i + 2, column=1).value
               duo_set.add(i)
               #duo_para = u'%s%s%s\%s' % ( duo_num,duanhao, duo_sheet.cell(row=i + 2, column=4).value ,huiche) # 问题
               print str
               duo_para = str(duo_num) + u'、' + duo_sheet.cell(row=i + 2, column=4).value + u'\n'  # 问题
               
               duo_temp = u'   A、%s\n   B、%s\n   C、%s\n   D、%s\n' % (
               duo_sheet.cell(row=i + 2, column=5).value, duo_sheet.cell(row=i + 2, column=6).value,
               duo_sheet.cell(row=i + 2, column=7).value, duo_sheet.cell(row=i + 2, column=8).value)
               duo_para += duo_temp

               if duo_sheet.cell(row=i + 2, column=9).value:
                   duo_para += u'   E、'+ duo_sheet.cell(row=i + 2, column=9).value + u'\n'
               print duo_para
               duo_paragraph = document.add_paragraph(duo_para)

               answer_duo_para += str(duo_num) + u'、' + duo_sheet.cell(row=i + 2, column=10).value + u'   '
               
               duo_num+=1
            else:
                pass
            
        answerdocument.add_paragraph(answer_duo_para) #答案

        panduanti_para = u'第三部分  判断题(10题)'
        panduan_paragraph = document.add_paragraph(panduanti_para)
        answerdocument.add_paragraph(panduanti_para) #答案
        
        panduan_num=1
        panduan_set = set()
        answer_panduan_para = u''
        while(len(panduan_set) < 10):
            i = random.randint(1, panduan_sheet.max_row-2)
            if i not in panduan_set:
               panduan_set.add(i)
               panduan_para = str(panduan_num) + u'、' + panduan_sheet.cell(row=i + 2, column=4).value + u'（ ）'  # 问题
               print dan_sheet.cell(row=i + 2, column=1).value
               
               print panduan_para
               panduan_paragraph = document.add_paragraph(panduan_para)

               answer_panduan_para += str(panduan_num) + u'、' + panduan_sheet.cell(row=i + 2, column=10).value + u'   '
               
               panduan_num+=1
            else:
                pass
            
        answerdocument.add_paragraph(answer_panduan_para) #答案
        
        # 保存文件
        length = len(filepath)
        for i in range(length - 1, -1, -1):
            if (filepath[i] == '/'):
                break
        savepath = u""
        for j in range(0, i + 1):
            savepath += filepath[j]

        #file_extension= u'%s %s' (count,'.docx')
        #u".docx"
        answer_savepath = savepath + u'测试'+ str(count) + u'答案.docx'
        savepath = savepath + u'测试'+ str(count) + u'.docx'
        
        print savepath
        document.save(savepath)
        answerdocument.save(answer_savepath)

button1=Button(fram,text="选择文件且执行程序",command=filefound, width=20, height=1).grid(row=0,column=0)
button2=Button(fram,text="退出", command=tk.quit,width=20, height=1).grid(row=2,column=0)
#print fram.size()
mainloop()
